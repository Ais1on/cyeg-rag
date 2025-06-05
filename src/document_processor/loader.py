"""
文档加载器 - 支持多种格式的威胁情报文档加载
"""
import os
from typing import List, Dict, Any, Optional
from pathlib import Path
import requests
from loguru import logger
import pypdf
import docx
from bs4 import BeautifulSoup
import json


class Document:
    """文档数据结构"""
    
    def __init__(self, content: str, metadata: Dict[str, Any] = None):
        self.content = content
        self.metadata = metadata or {}
    
    def __repr__(self):
        return f"Document(content_length={len(self.content)}, metadata={self.metadata})"


class DocumentLoader:
    """文档加载器"""
    
    def __init__(self):
        self.supported_formats = ['.pdf', '.docx', '.txt', '.html', '.json']
    
    def load_file(self, file_path: str) -> List[Document]:
        """加载单个文件"""
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")
        
        extension = file_path.suffix.lower()
        
        if extension not in self.supported_formats:
            raise ValueError(f"不支持的文件格式: {extension}")
        
        logger.info(f"加载文件: {file_path}")
        
        try:
            if extension == '.pdf':
                return self._load_pdf(file_path)
            elif extension == '.docx':
                return self._load_docx(file_path)
            elif extension == '.txt':
                return self._load_txt(file_path)
            elif extension == '.html':
                return self._load_html(file_path)
            elif extension == '.json':
                return self._load_json(file_path)
        except Exception as e:
            logger.error(f"加载文件失败 {file_path}: {str(e)}")
            raise
    
    def load_directory(self, directory_path: str, recursive: bool = True) -> List[Document]:
        """加载目录中的所有文档"""
        directory_path = Path(directory_path)
        
        if not directory_path.exists():
            raise FileNotFoundError(f"目录不存在: {directory_path}")
        
        documents = []
        pattern = "**/*" if recursive else "*"
        
        for file_path in directory_path.glob(pattern):
            if file_path.is_file() and file_path.suffix.lower() in self.supported_formats:
                try:
                    docs = self.load_file(file_path)
                    documents.extend(docs)
                except Exception as e:
                    logger.warning(f"跳过文件 {file_path}: {str(e)}")
        
        logger.info(f"从目录 {directory_path} 加载了 {len(documents)} 个文档")
        return documents
    
    def load_from_url(self, url: str) -> List[Document]:
        """从URL加载文档"""
        try:
            response = requests.get(url, timeout=30)
            response.raise_for_status()
            
            content_type = response.headers.get('content-type', '').lower()
            
            metadata = {
                'source': url,
                'content_type': content_type,
                'status_code': response.status_code
            }
            
            if 'application/pdf' in content_type:
                return self._parse_pdf_content(response.content, metadata)
            elif 'text/html' in content_type:
                return self._parse_html_content(response.text, metadata)
            elif 'application/json' in content_type:
                return self._parse_json_content(response.text, metadata)
            else:
                # 默认按文本处理
                content = response.text
                return [Document(content=content, metadata=metadata)]
                
        except Exception as e:
            logger.error(f"从URL加载失败 {url}: {str(e)}")
            raise
    
    def _load_pdf(self, file_path: Path) -> List[Document]:
        """加载PDF文件"""
        documents = []
        
        with open(file_path, 'rb') as file:
            pdf_reader = pypdf.PdfReader(file)
            
            for page_num, page in enumerate(pdf_reader.pages):
                text = page.extract_text()
                if text.strip():
                    metadata = {
                        'source': str(file_path),
                        'page': page_num + 1,
                        'total_pages': len(pdf_reader.pages),
                        'file_type': 'pdf'
                    }
                    documents.append(Document(content=text, metadata=metadata))
        
        return documents
    
    def _load_docx(self, file_path: Path) -> List[Document]:
        """加载DOCX文件"""
        doc = docx.Document(file_path)
        text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
        
        metadata = {
            'source': str(file_path),
            'file_type': 'docx',
            'paragraphs': len(doc.paragraphs)
        }
        
        return [Document(content=text, metadata=metadata)]
    
    def _load_txt(self, file_path: Path) -> List[Document]:
        """加载TXT文件"""
        with open(file_path, 'r', encoding='utf-8') as file:
            content = file.read()
        
        metadata = {
            'source': str(file_path),
            'file_type': 'txt',
            'encoding': 'utf-8'
        }
        
        return [Document(content=content, metadata=metadata)]
    
    def _load_html(self, file_path: Path) -> List[Document]:
        """加载HTML文件"""
        with open(file_path, 'r', encoding='utf-8') as file:
            content = file.read()
        
        return self._parse_html_content(content, {'source': str(file_path), 'file_type': 'html'})
    
    def _load_json(self, file_path: Path) -> List[Document]:
        """加载JSON文件（威胁情报格式）"""
        with open(file_path, 'r', encoding='utf-8') as file:
            data = json.load(file)
        
        return self._parse_json_content(json.dumps(data), {'source': str(file_path), 'file_type': 'json'})
    
    def _parse_html_content(self, html_content: str, base_metadata: Dict) -> List[Document]:
        """解析HTML内容"""
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # 移除脚本和样式标签
        for script in soup(["script", "style"]):
            script.extract()
        
        text = soup.get_text()
        # 清理文本
        lines = (line.strip() for line in text.splitlines())
        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
        text = ' '.join(chunk for chunk in chunks if chunk)
        
        metadata = {**base_metadata, 'title': soup.title.string if soup.title else ''}
        
        return [Document(content=text, metadata=metadata)]
    
    def _parse_json_content(self, json_content: str, base_metadata: Dict) -> List[Document]:
        """解析JSON内容（威胁情报专用）"""
        try:
            data = json.loads(json_content)
            documents = []
            
            # 如果是威胁情报标准格式（如STIX）
            if isinstance(data, dict):
                if 'objects' in data:  # STIX格式
                    for obj in data['objects']:
                        content = json.dumps(obj, ensure_ascii=False, indent=2)
                        metadata = {
                            **base_metadata,
                            'object_type': obj.get('type', 'unknown'),
                            'object_id': obj.get('id', ''),
                            'stix_format': True
                        }
                        documents.append(Document(content=content, metadata=metadata))
                else:
                    # 普通JSON
                    content = json.dumps(data, ensure_ascii=False, indent=2)
                    documents.append(Document(content=content, metadata=base_metadata))
            elif isinstance(data, list):
                for i, item in enumerate(data):
                    content = json.dumps(item, ensure_ascii=False, indent=2)
                    metadata = {**base_metadata, 'item_index': i}
                    documents.append(Document(content=content, metadata=metadata))
            
            return documents
            
        except json.JSONDecodeError as e:
            logger.error(f"JSON解析错误: {str(e)}")
            # 如果JSON解析失败，作为普通文本处理
            return [Document(content=json_content, metadata=base_metadata)]
    
    def _parse_pdf_content(self, pdf_content: bytes, base_metadata: Dict) -> List[Document]:
        """解析PDF二进制内容"""
        import io
        documents = []
        
        pdf_reader = pypdf.PdfReader(io.BytesIO(pdf_content))
        
        for page_num, page in enumerate(pdf_reader.pages):
            text = page.extract_text()
            if text.strip():
                metadata = {
                    **base_metadata,
                    'page': page_num + 1,
                    'total_pages': len(pdf_reader.pages),
                    'file_type': 'pdf'
                }
                documents.append(Document(content=text, metadata=metadata))
        
        return documents 